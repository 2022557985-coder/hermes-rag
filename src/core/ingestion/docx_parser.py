"""DOCX parser using python-docx."""

from pathlib import Path
from typing import Any

from .parser_factory import BaseParser


class DocxParser(BaseParser):
    """Parse DOCX documents."""

    def parse(self, source: str) -> dict[str, Any]:
        """Parse a DOCX file.

        Args:
            source: Path to DOCX file.

        Returns:
            dict with text, tables, metadata.
        """
        from docx import Document

        doc = Document(source)
        full_text = []
        tables = []
        metadata = {
            "source": Path(source).name,
            "type": "docx",
            "total_paragraphs": len(doc.paragraphs),
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect heading level
                if para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split()[-1])
                        prefix = "#" * level + " "
                        full_text.append(prefix + text)
                    except ValueError:
                        full_text.append(text)
                else:
                    full_text.append(text)

        # Extract tables
        for i, table in enumerate(doc.tables):
            md_table = self._table_to_markdown(table)
            if md_table:
                tables.append({
                    "index": i,
                    "content": md_table,
                })

        return {
            "text": "\n\n".join(full_text),
            "tables": tables,
            "metadata": metadata,
        }